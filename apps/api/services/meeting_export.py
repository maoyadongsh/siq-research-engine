"""Private, version-aware meeting exports and one-time download tickets."""

from __future__ import annotations

import hashlib
import html
import json
import os
import re
import secrets
import unicodedata
from dataclasses import dataclass
from datetime import datetime, timedelta, timezone
from io import BytesIO
from pathlib import Path
from typing import Any
from uuid import uuid4

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from sqlalchemy import and_, or_, update
from sqlmodel import col, select
from sqlmodel.ext.asyncio.session import AsyncSession

from services.meeting_contracts import (
    ArtifactState,
    ArtifactType,
    MeetingArtifact,
    MeetingExportContent,
    MeetingExportCreateRequest,
    MeetingExportFormat,
    MeetingExportResponse,
    MeetingJob,
    MeetingJobKind,
    MeetingJobState,
    MeetingSession,
    MeetingStreamTicket,
    TranscriptSegmentResponse,
    utcnow,
)
from services.meeting_event_store import MeetingEventStore, decode_json, encode_json
from services.meeting_repository import (
    MeetingInvalidOperation,
    MeetingRepository,
    MeetingResourceNotFound,
    export_response,
)
from services.path_config import BACKEND_DATA_ROOT

_ID_RE = re.compile(r"^[A-Za-z0-9_-]{1,64}$")
_MARKDOWN_RE = re.compile(r"([\\`*_[\]{}()#+.!|>~-])")
_CONTROL_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")
_FILENAME_RE = re.compile(r"[^\w\-.() ]+", re.UNICODE)
_URL_RE = re.compile(r"https?://[^\s]+", re.IGNORECASE)
_SECRET_RE = re.compile(r"(?i)(authorization|api[_-]?key|token|secret|password)\s*[:=]\s*[^\s,;]+")

_EXTENSIONS = {
    MeetingExportFormat.MARKDOWN.value: "md",
    MeetingExportFormat.TXT.value: "txt",
    MeetingExportFormat.SRT.value: "srt",
    MeetingExportFormat.VTT.value: "vtt",
    MeetingExportFormat.JSON.value: "json",
    MeetingExportFormat.DOCX.value: "docx",
    MeetingExportFormat.PDF.value: "pdf",
}
_MEDIA_TYPES = {
    MeetingExportFormat.MARKDOWN.value: "text/markdown; charset=utf-8",
    MeetingExportFormat.TXT.value: "text/plain; charset=utf-8",
    MeetingExportFormat.SRT.value: "application/x-subrip; charset=utf-8",
    MeetingExportFormat.VTT.value: "text/vtt; charset=utf-8",
    MeetingExportFormat.JSON.value: "application/json; charset=utf-8",
    MeetingExportFormat.DOCX.value: "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    MeetingExportFormat.PDF.value: "application/pdf",
}
_MINUTES_SECTIONS = (
    ("agenda_topics", "议题"),
    ("chapters", "章节"),
    ("decisions", "决定"),
    ("open_questions", "待确认问题"),
    ("risks", "风险"),
    ("action_items", "待办"),
    ("speaker_viewpoints", "发言人观点"),
)
_DOCX_BODY_FONT = "Microsoft YaHei"
_DOCX_HEADING_FONT = "Microsoft YaHei"


class MeetingExportError(RuntimeError):
    def __init__(
        self,
        code: str,
        message: str,
        *,
        retryable: bool = False,
    ) -> None:
        super().__init__(message)
        self.code = code
        self.retryable = retryable


@dataclass(frozen=True)
class MeetingExportSettings:
    max_bytes: int = 20 * 1024 * 1024
    max_segments: int = 200_000
    ticket_ttl_seconds: int = 120
    lease_seconds: int = 120
    retry_delay_seconds: int = 20

    @classmethod
    def from_env(cls) -> "MeetingExportSettings":
        def integer(name: str, default: int, minimum: int, maximum: int) -> int:
            try:
                value = int(os.getenv(name, str(default)))
            except ValueError as exc:
                raise MeetingExportError("EXPORT_CONFIGURATION_INVALID", f"{name} must be an integer") from exc
            if value < minimum or value > maximum:
                raise MeetingExportError(
                    "EXPORT_CONFIGURATION_INVALID",
                    f"{name} must be between {minimum} and {maximum}",
                )
            return value

        return cls(
            max_bytes=integer(
                "SIQ_MEETING_EXPORT_MAX_BYTES",
                20 * 1024 * 1024,
                64 * 1024,
                200 * 1024 * 1024,
            ),
            max_segments=integer("SIQ_MEETING_EXPORT_MAX_SEGMENTS", 200_000, 1, 1_000_000),
            ticket_ttl_seconds=integer("SIQ_MEETING_EXPORT_TICKET_TTL_SECONDS", 120, 30, 900),
            lease_seconds=integer("SIQ_MEETING_EXPORT_LEASE_SECONDS", 120, 30, 3_600),
            retry_delay_seconds=integer("SIQ_MEETING_EXPORT_RETRY_DELAY_SECONDS", 20, 1, 3_600),
        )


@dataclass(frozen=True)
class RenderedMeetingExport:
    payload: bytes
    filename: str
    media_type: str
    transcript_revision: int
    revision_vector: dict[str, int]
    source_artifact_id: str | None = None
    source_artifact_version: int | None = None


@dataclass(frozen=True)
class PersistedMeetingExport:
    storage_key: str
    sha256: str
    byte_size: int
    filename: str
    media_type: str


@dataclass(frozen=True)
class DownloadableMeetingExport:
    path: Path
    filename: str
    media_type: str
    byte_size: int
    sha256: str


def _aware(value: datetime) -> datetime:
    return value if value.tzinfo is not None else value.replace(tzinfo=timezone.utc)


def _diagnostic(exc: BaseException) -> str:
    value = f"{type(exc).__name__}: {exc}"
    value = _URL_RE.sub("[redacted-url]", value)
    value = _SECRET_RE.sub(r"\1=[redacted]", value)
    return value[:1000]


def _safe_component(value: str) -> str:
    if not _ID_RE.fullmatch(value):
        raise MeetingExportError("EXPORT_STORAGE_ID_INVALID", "export identifier is invalid")
    return value


def _safe_filename(title: str, export_id: str, export_format: str) -> str:
    normalized = unicodedata.normalize("NFKC", title or "meeting")
    normalized = normalized.splitlines()[0] if normalized.splitlines() else "meeting"
    normalized = _CONTROL_RE.sub("", normalized).replace("/", " ").replace("\\", " ")
    normalized = normalized.replace("\r", " ").replace("\n", " ")
    normalized = _FILENAME_RE.sub("_", normalized)
    normalized = re.sub(r"\s+", "-", normalized).strip(" .-_")[:80] or "meeting"
    suffix = _EXTENSIONS[export_format]
    return f"{normalized}-{export_id[:8]}.{suffix}"


def _plain_text(value: Any) -> str:
    text = unicodedata.normalize("NFC", str(value or ""))
    text = _CONTROL_RE.sub("", text)
    return text.replace("\r\n", "\n").replace("\r", "\n")


def _single_line(value: Any) -> str:
    return re.sub(r"\s+", " ", _plain_text(value)).strip().replace("-->", "->")


def _docx_text(value: Any) -> str:
    """Return literal XML 1.0 text; callers never interpret markup or fields."""

    normalized = unicodedata.normalize("NFC", str(value or ""))
    normalized = normalized.replace("\r\n", "\n").replace("\r", "\n")
    normalized = _CONTROL_RE.sub("", normalized)
    return "".join(
        character
        for character in normalized
        if character in {"\t", "\n"}
        or 0x20 <= ord(character) <= 0xD7FF
        or 0xE000 <= ord(character) <= 0xFFFD
        or 0x10000 <= ord(character) <= 0x10FFFF
    )


def _docx_set_font(target: Any, name: str, size: float | None = None) -> None:
    target.font.name = name
    if size is not None:
        target.font.size = Pt(size)
    properties = target._element.get_or_add_rPr()
    properties.rFonts.set(qn("w:eastAsia"), name)


def _docx_add_run(
    paragraph: Any,
    value: Any,
    *,
    bold: bool = False,
    italic: bool = False,
    size: float | None = None,
) -> Any:
    run = paragraph.add_run(_docx_text(value))
    _docx_set_font(run, _DOCX_BODY_FONT, size)
    run.bold = bold
    run.italic = italic
    return run


def _meeting_date(meeting: MeetingSession) -> str:
    value = meeting.started_at or meeting.created_at
    return _aware(value).isoformat(timespec="minutes") if value else "未记录"


def _new_docx(meeting: MeetingSession, document_kind: str, transcript_source: str) -> Any:
    document = Document()
    for section in document.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)

    style_fonts = {
        "Normal": (_DOCX_BODY_FONT, 10.5),
        "Body Text": (_DOCX_BODY_FONT, 10.5),
        "List Bullet": (_DOCX_BODY_FONT, 10.5),
        "Title": (_DOCX_HEADING_FONT, 20),
        "Heading 1": (_DOCX_HEADING_FONT, 15),
        "Heading 2": (_DOCX_HEADING_FONT, 12),
        "Heading 3": (_DOCX_HEADING_FONT, 11),
    }
    for style_name, (font_name, font_size) in style_fonts.items():
        try:
            style = document.styles[style_name]
        except KeyError:
            continue
        _docx_set_font(style, font_name, font_size)

    title = document.add_heading(_docx_text(_single_line(meeting.title)), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _docx_add_run(subtitle, document_kind, bold=True, size=12)

    metadata = document.add_paragraph()
    _docx_add_run(metadata, "会议日期：", bold=True)
    _docx_add_run(metadata, _meeting_date(meeting))
    _docx_add_run(metadata, "\n会议 ID：", bold=True)
    _docx_add_run(metadata, meeting.id)
    _docx_add_run(metadata, "\n文字来源：", bold=True)
    _docx_add_run(metadata, "当前显示文字" if transcript_source == "display" else "ASR 原文")

    properties = document.core_properties
    properties.title = _docx_text(_single_line(meeting.title))[:255]
    properties.subject = _docx_text(document_kind)[:255]
    properties.category = "SIQ Meeting Export"
    properties.language = _docx_text(meeting.language or "zh-CN")[:32]
    return document


def _docx_bytes(document: Any) -> bytes:
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _markdown(value: Any) -> str:
    escaped = html.escape(_plain_text(value), quote=False)
    return _MARKDOWN_RE.sub(r"\\\1", escaped).replace("\n", "  \n")


def _timestamp(milliseconds: int, *, subtitle: str | None = None) -> str:
    value = max(0, int(milliseconds))
    hours, remainder = divmod(value, 3_600_000)
    minutes, remainder = divmod(remainder, 60_000)
    seconds, millis = divmod(remainder, 1_000)
    separator = "," if subtitle == "srt" else "."
    return f"{hours:02d}:{minutes:02d}:{seconds:02d}{separator}{millis:03d}"


def _segment_text(segment: TranscriptSegmentResponse, source: str) -> str:
    return segment.display_text if source == "display" else segment.raw_text


def _speaker(segment: TranscriptSegmentResponse) -> str:
    return _single_line(segment.speaker_label or "发言人")


def _transcript_json(
    meeting: MeetingSession,
    segments: list[TranscriptSegmentResponse],
    source: str,
) -> dict[str, Any]:
    return {
        "schema_version": "siq.meeting.transcript_export.v1",
        "meeting": {
            "id": meeting.id,
            "title": meeting.title,
            "language": meeting.language,
            "started_at": meeting.started_at.isoformat() if meeting.started_at else None,
            "stopped_at": meeting.stopped_at.isoformat() if meeting.stopped_at else None,
        },
        "transcript_source": source,
        "segments": [
            {
                "segment_id": segment.id,
                "ordinal": segment.ordinal,
                "start_ms": segment.start_ms,
                "end_ms": segment.end_ms,
                "start_timestamp": _timestamp(segment.start_ms),
                "end_timestamp": _timestamp(segment.end_ms),
                "speaker": segment.speaker_label,
                "speaker_track_id": segment.speaker_track_id,
                "text": _segment_text(segment, source),
                "revision_no": segment.current_revision_no,
                "display_layer": segment.display_layer,
                "human_locked": segment.human_locked,
            }
            for segment in segments
        ],
    }


def _render_transcript(
    meeting: MeetingSession,
    segments: list[TranscriptSegmentResponse],
    export_format: str,
    source: str,
) -> bytes:
    if export_format == MeetingExportFormat.DOCX.value:
        document = _new_docx(meeting, "会议逐字稿", source)
        document.add_heading("逐字稿", level=1)
        if not segments:
            document.add_paragraph("本次会议暂无稳定逐字稿。")
        for segment in segments:
            header = document.add_paragraph()
            header.paragraph_format.keep_with_next = True
            _docx_add_run(
                header,
                f"[{_timestamp(segment.start_ms)} - {_timestamp(segment.end_ms)}] ",
                bold=True,
            )
            _docx_add_run(header, _speaker(segment), bold=True)
            body = document.add_paragraph()
            _docx_add_run(body, _segment_text(segment, source))
            evidence = document.add_paragraph()
            _docx_add_run(
                evidence,
                f"来源片段：{segment.id} · 修订版本：{segment.current_revision_no}",
                italic=True,
                size=8.5,
            )
        return _docx_bytes(document)
    if export_format == MeetingExportFormat.JSON.value:
        return (
            json.dumps(
                _transcript_json(meeting, segments, source),
                ensure_ascii=False,
                indent=2,
            )
            + "\n"
        ).encode("utf-8")
    if export_format == MeetingExportFormat.VTT.value:
        lines = ["WEBVTT", ""]
        for segment in segments:
            lines.extend(
                [
                    segment.id,
                    f"{_timestamp(segment.start_ms)} --> {_timestamp(segment.end_ms)}",
                    f"{_speaker(segment)}: {_single_line(_segment_text(segment, source))}",
                    "",
                ]
            )
        return "\n".join(lines).encode("utf-8")
    if export_format == MeetingExportFormat.SRT.value:
        lines: list[str] = []
        for index, segment in enumerate(segments, start=1):
            lines.extend(
                [
                    str(index),
                    f"{_timestamp(segment.start_ms, subtitle='srt')} --> {_timestamp(segment.end_ms, subtitle='srt')}",
                    f"{_speaker(segment)}: {_single_line(_segment_text(segment, source))}",
                    "",
                ]
            )
        return ("\n".join(lines) if lines else "\n").encode("utf-8")
    if export_format == MeetingExportFormat.MARKDOWN.value:
        lines = [
            f"# {_markdown(meeting.title)}",
            "",
            f"- 会议 ID: `{meeting.id}`",
            f"- 逐字稿层: `{source}`",
            "",
            "## 逐字稿",
            "",
        ]
        for segment in segments:
            lines.extend(
                [
                    f"### {_timestamp(segment.start_ms)} - {_markdown(_speaker(segment))}",
                    "",
                    _markdown(_segment_text(segment, source)),
                    "",
                    f"证据: `{segment.id}` / {_timestamp(segment.start_ms)} - {_timestamp(segment.end_ms)}",
                    "",
                ]
            )
        return "\n".join(lines).encode("utf-8")
    lines = [
        _plain_text(meeting.title),
        f"Meeting ID: {meeting.id}",
        f"Transcript source: {source}",
        "",
    ]
    for segment in segments:
        lines.append(
            f"[{_timestamp(segment.start_ms)} - {_timestamp(segment.end_ms)}] "
            f"{_speaker(segment)}: {_single_line(_segment_text(segment, source))} "
            f"(segment {segment.id})"
        )
    return ("\n".join(lines) + "\n").encode("utf-8")


def _evidence_ids(payload: dict[str, Any]) -> list[str]:
    values: list[str] = []
    for key, _ in _MINUTES_SECTIONS:
        items = payload.get(key, [])
        if not isinstance(items, list):
            raise MeetingExportError("EXPORT_ARTIFACT_SCHEMA_INVALID", "minutes artifact has an invalid section")
        for item in items:
            if not isinstance(item, dict):
                raise MeetingExportError("EXPORT_ARTIFACT_SCHEMA_INVALID", "minutes artifact item is invalid")
            source_ids = item.get("source_segment_ids")
            if not isinstance(source_ids, list) or not source_ids:
                raise MeetingExportError("EXPORT_ARTIFACT_SCHEMA_INVALID", "minutes evidence is missing")
            values.extend(str(value) for value in source_ids)
    return list(dict.fromkeys(values))


def _render_minutes(
    meeting: MeetingSession,
    artifact: MeetingArtifact,
    segments: list[TranscriptSegmentResponse],
    export_format: str,
    source: str,
) -> bytes:
    payload = decode_json(artifact.content_json, None)
    if not isinstance(payload, dict):
        raise MeetingExportError("EXPORT_ARTIFACT_SCHEMA_INVALID", "minutes artifact is not structured JSON")
    by_id = {segment.id: segment for segment in segments}
    evidence_ids = _evidence_ids(payload)
    if any(segment_id not in by_id for segment_id in evidence_ids):
        raise MeetingExportError("EXPORT_ARTIFACT_SCHEMA_INVALID", "minutes artifact cites unknown evidence")
    evidence = [
        {
            "segment_id": segment_id,
            "start_ms": by_id[segment_id].start_ms,
            "end_ms": by_id[segment_id].end_ms,
            "start_timestamp": _timestamp(by_id[segment_id].start_ms),
            "end_timestamp": _timestamp(by_id[segment_id].end_ms),
            "speaker": by_id[segment_id].speaker_label,
            "text": _segment_text(by_id[segment_id], source),
            "revision_no": by_id[segment_id].current_revision_no,
        }
        for segment_id in evidence_ids
    ]
    if export_format == MeetingExportFormat.DOCX.value:
        document = _new_docx(meeting, "会议纪要", source)
        version = document.add_paragraph()
        _docx_add_run(version, "纪要版本：", bold=True)
        _docx_add_run(version, str(artifact.version))
        _docx_add_run(version, "\n产物 ID：", bold=True)
        _docx_add_run(version, artifact.id)
        _docx_add_run(version, "\n产物状态：", bold=True)
        _docx_add_run(version, artifact.state)

        document.add_heading("概览", level=1)
        overview = document.add_paragraph()
        _docx_add_run(overview, payload.get("overview", ""))
        for key, title in _MINUTES_SECTIONS:
            items = payload.get(key, [])
            if not items:
                continue
            document.add_heading(_docx_text(title), level=1)
            for item in items:
                paragraph = document.add_paragraph(style="List Bullet")
                _docx_add_run(paragraph, item.get("text", ""))
                citations: list[str] = []
                for raw_segment_id in item["source_segment_ids"]:
                    segment = by_id[str(raw_segment_id)]
                    citations.append(
                        f"{_timestamp(segment.start_ms)} - {_timestamp(segment.end_ms)} ({segment.id})"
                    )
                evidence_paragraph = document.add_paragraph()
                evidence_paragraph.paragraph_format.left_indent = Cm(0.75)
                _docx_add_run(
                    evidence_paragraph,
                    f"证据：{', '.join(citations)}",
                    italic=True,
                    size=8.5,
                )

        if evidence:
            document.add_heading("证据索引", level=1)
            for item in evidence:
                paragraph = document.add_paragraph(style="List Bullet")
                _docx_add_run(
                    paragraph,
                    f"[{item['start_timestamp']} - {item['end_timestamp']}] ",
                    bold=True,
                )
                _docx_add_run(paragraph, item.get("speaker") or "发言人", bold=True)
                _docx_add_run(paragraph, f"：{item['text']}")
                detail = document.add_paragraph()
                detail.paragraph_format.left_indent = Cm(0.75)
                _docx_add_run(
                    detail,
                    f"来源片段：{item['segment_id']} · 修订版本：{item['revision_no']}",
                    italic=True,
                    size=8.5,
                )
        return _docx_bytes(document)
    if export_format == MeetingExportFormat.JSON.value:
        value = {
            "schema_version": "siq.meeting.minutes_export.v1",
            "meeting": {"id": meeting.id, "title": meeting.title},
            "artifact": {
                "id": artifact.id,
                "artifact_type": artifact.artifact_type,
                "version": artifact.version,
                "state": artifact.state,
                "transcript_revision": artifact.transcript_revision,
            },
            "minutes": payload,
            "evidence": evidence,
        }
        return (json.dumps(value, ensure_ascii=False, indent=2) + "\n").encode("utf-8")

    lines = [
        f"# {_markdown(meeting.title)}",
        "",
        f"- 纪要版本: `{artifact.version}`",
        f"- 产物 ID: `{artifact.id}`",
        f"- 产物状态: `{artifact.state}`",
        "",
        "## 概览",
        "",
        _markdown(payload.get("overview", "")),
    ]
    for key, title in _MINUTES_SECTIONS:
        items = payload.get(key, [])
        if not items:
            continue
        lines.extend(["", f"## {title}", ""])
        for item in items:
            citations = []
            for segment_id in item["source_segment_ids"]:
                segment = by_id[str(segment_id)]
                citations.append(f"{_timestamp(segment.start_ms)} (`{segment.id}`)")
            lines.append(f"- {_markdown(item.get('text', ''))}  ")
            lines.append(f"  证据: {', '.join(citations)}")
    if evidence:
        lines.extend(["", "## 证据索引", ""])
        for item in evidence:
            lines.append(
                f"- {_timestamp(item['start_ms'])} - `{item['segment_id']}` - "
                f"{_markdown(item.get('speaker') or '发言人')}: {_markdown(item['text'])}"
            )
    return ("\n".join(lines).strip() + "\n").encode("utf-8")


class MeetingExportStorage:
    def __init__(self, root: Path | None = None, *, max_bytes: int) -> None:
        configured = os.getenv("SIQ_MEETING_EXPORT_ROOT", "").strip()
        self.root = (
            (root or (Path(configured) if configured else BACKEND_DATA_ROOT / "meeting_exports")).expanduser().resolve()
        )
        self.max_bytes = max_bytes
        self.root.mkdir(parents=True, exist_ok=True, mode=0o700)

    def _target(
        self,
        owner_user_id: int,
        meeting_id: str,
        export_id: str,
        extension: str,
    ) -> Path:
        owner = _safe_component(str(owner_user_id))
        meeting = _safe_component(meeting_id)
        export = _safe_component(export_id)
        target = (self.root / owner / meeting / "exports" / f"{export}.{extension}").resolve()
        if self.root not in target.parents:
            raise MeetingExportError("EXPORT_STORAGE_PATH_INVALID", "export path escaped its root")
        return target

    def persist(
        self,
        owner_user_id: int,
        meeting_id: str,
        export_id: str,
        rendered: RenderedMeetingExport,
    ) -> PersistedMeetingExport:
        if not rendered.payload or len(rendered.payload) > self.max_bytes:
            raise MeetingExportError("EXPORT_TOO_LARGE", "export exceeds its byte limit")
        extension = rendered.filename.rsplit(".", 1)[-1]
        target = self._target(owner_user_id, meeting_id, export_id, extension)
        target.parent.mkdir(parents=True, exist_ok=True, mode=0o700)
        temporary = target.with_name(f".{target.name}.{uuid4().hex}.tmp")
        digest = hashlib.sha256(rendered.payload).hexdigest()
        try:
            descriptor = os.open(temporary, os.O_CREAT | os.O_EXCL | os.O_WRONLY, 0o600)
            with os.fdopen(descriptor, "wb") as output:
                output.write(rendered.payload)
                output.flush()
                os.fsync(output.fileno())
            os.replace(temporary, target)
        except OSError as exc:
            temporary.unlink(missing_ok=True)
            raise MeetingExportError(
                "EXPORT_STORAGE_UNAVAILABLE",
                "export storage is unavailable",
                retryable=True,
            ) from exc
        return PersistedMeetingExport(
            storage_key=target.relative_to(self.root).as_posix(),
            sha256=digest,
            byte_size=len(rendered.payload),
            filename=rendered.filename,
            media_type=rendered.media_type,
        )

    def resolve(self, metadata: dict[str, Any]) -> DownloadableMeetingExport:
        storage_key = str(metadata.get("storage_key") or "")
        if not storage_key or Path(storage_key).is_absolute():
            raise MeetingExportError("EXPORT_FILE_UNAVAILABLE", "export file is unavailable")
        path = (self.root / storage_key).resolve()
        if self.root not in path.parents:
            raise MeetingExportError("EXPORT_FILE_UNAVAILABLE", "export file is unavailable")
        filename = str(metadata.get("filename") or "")
        media_type = str(metadata.get("media_type") or "application/octet-stream")
        sha256 = str(metadata.get("sha256") or "")
        byte_size = int(metadata.get("byte_size") or 0)
        if not filename or "\r" in filename or "\n" in filename or "/" in filename or "\\" in filename:
            raise MeetingExportError("EXPORT_FILE_UNAVAILABLE", "export filename is invalid")
        try:
            actual_size = path.stat().st_size
            digest = hashlib.sha256()
            with path.open("rb") as source:
                for block in iter(lambda: source.read(1024 * 1024), b""):
                    digest.update(block)
        except OSError as exc:
            raise MeetingExportError("EXPORT_FILE_UNAVAILABLE", "export file is unavailable") from exc
        if actual_size != byte_size or digest.hexdigest() != sha256:
            raise MeetingExportError("EXPORT_FILE_INTEGRITY_FAILED", "export file integrity failed")
        return DownloadableMeetingExport(path, filename, media_type, byte_size, sha256)


class MeetingExportService:
    def __init__(
        self,
        session: AsyncSession,
        *,
        settings: MeetingExportSettings | None = None,
        storage: MeetingExportStorage | None = None,
    ) -> None:
        self.session = session
        self.settings = settings or MeetingExportSettings.from_env()
        self.storage = storage or MeetingExportStorage(max_bytes=self.settings.max_bytes)
        self.repository = MeetingRepository(session)
        self.events = MeetingEventStore(session)

    async def process_export(self, export_id: str, owner_user_id: int, worker_id: str) -> None:
        artifact, job = await self.repository.get_export(
            (await self._owned_export_artifact(export_id, owner_user_id)).meeting_id,
            export_id,
            owner_user_id,
        )
        if artifact.state == ArtifactState.READY.value and job.state == MeetingJobState.SUCCEEDED.value:
            return
        claimed = await self._claim_exact(job.id, worker_id)
        if not claimed:
            return
        try:
            meeting, request = await self._load_request(job.id, worker_id)
            rendered = await self._render(meeting, artifact, job, request, owner_user_id)
            persisted = self.storage.persist(
                owner_user_id,
                meeting.id,
                artifact.id,
                rendered,
            )
            await self._complete(job.id, worker_id, rendered, persisted)
        except Exception as exc:
            await self._fail(job.id, worker_id, exc)

    async def claim_next(self, worker_id: str) -> str | None:
        now = utcnow()
        retry_before = now - timedelta(seconds=self.settings.retry_delay_seconds)
        eligible = and_(
            MeetingJob.job_kind == MeetingJobKind.EXPORT.value,
            MeetingJob.attempt < MeetingJob.max_attempts,
            or_(
                MeetingJob.state == MeetingJobState.QUEUED.value,
                and_(
                    MeetingJob.state == MeetingJobState.RETRY_WAIT.value,
                    MeetingJob.updated_at <= retry_before,
                ),
                and_(
                    col(MeetingJob.state).in_([MeetingJobState.LEASED.value, MeetingJobState.RUNNING.value]),
                    MeetingJob.lease_until.is_not(None),
                    MeetingJob.lease_until < now,
                ),
            ),
        )
        candidate = (
            select(MeetingJob.id)
            .where(eligible)
            .order_by(MeetingJob.created_at, MeetingJob.id)
            .limit(1)
            .scalar_subquery()
        )
        result = await self.session.exec(
            update(MeetingJob)
            .where(MeetingJob.id == candidate, eligible)
            .values(
                state=MeetingJobState.LEASED.value,
                attempt=MeetingJob.attempt + 1,
                lease_owner=worker_id[:100],
                lease_until=now + timedelta(seconds=self.settings.lease_seconds),
                public_error_code=None,
                internal_diagnostic=None,
                updated_at=now,
            )
            .returning(MeetingJob.id)
        )
        job_id = result.scalar_one_or_none()
        await self.session.commit()
        return job_id

    async def process_claimed(self, job_id: str, worker_id: str) -> None:
        job = await self.session.get(MeetingJob, job_id)
        if job is None:
            return
        artifact_id = str(decode_json(job.input_json, {}).get("export_id") or "")
        if not artifact_id:
            await self._fail(
                job_id,
                worker_id,
                MeetingExportError("EXPORT_INPUT_INVALID", "export job input is invalid"),
            )
            return
        meeting = await self.session.get(MeetingSession, job.meeting_id)
        if meeting is None:
            await self._fail(
                job_id,
                worker_id,
                MeetingExportError("EXPORT_MEETING_UNAVAILABLE", "meeting is unavailable"),
            )
            return
        try:
            artifact = await self._owned_export_artifact(artifact_id, meeting.owner_user_id)
            request_meeting, request = await self._load_request(job_id, worker_id)
            rendered = await self._render(
                request_meeting,
                artifact,
                job,
                request,
                meeting.owner_user_id,
            )
            persisted = self.storage.persist(
                meeting.owner_user_id,
                meeting.id,
                artifact.id,
                rendered,
            )
            await self._complete(job_id, worker_id, rendered, persisted)
        except Exception as exc:
            await self._fail(job_id, worker_id, exc)

    async def _claim_exact(self, job_id: str, worker_id: str) -> bool:
        now = utcnow()
        retry_before = now - timedelta(seconds=self.settings.retry_delay_seconds)
        eligible = or_(
            MeetingJob.state == MeetingJobState.QUEUED.value,
            and_(
                MeetingJob.state == MeetingJobState.RETRY_WAIT.value,
                MeetingJob.updated_at <= retry_before,
            ),
            and_(
                col(MeetingJob.state).in_([MeetingJobState.LEASED.value, MeetingJobState.RUNNING.value]),
                MeetingJob.lease_until.is_not(None),
                MeetingJob.lease_until < now,
            ),
        )
        result = await self.session.exec(
            update(MeetingJob)
            .where(
                MeetingJob.id == job_id,
                MeetingJob.job_kind == MeetingJobKind.EXPORT.value,
                MeetingJob.attempt < MeetingJob.max_attempts,
                eligible,
            )
            .values(
                state=MeetingJobState.LEASED.value,
                attempt=MeetingJob.attempt + 1,
                lease_owner=worker_id[:100],
                lease_until=now + timedelta(seconds=self.settings.lease_seconds),
                public_error_code=None,
                internal_diagnostic=None,
                updated_at=now,
            )
        )
        await self.session.commit()
        return bool(result.rowcount)

    async def _load_request(self, job_id: str, worker_id: str) -> tuple[MeetingSession, MeetingExportCreateRequest]:
        job = (
            await self.session.exec(
                select(MeetingJob).where(
                    MeetingJob.id == job_id,
                    MeetingJob.job_kind == MeetingJobKind.EXPORT.value,
                    MeetingJob.state == MeetingJobState.LEASED.value,
                    MeetingJob.lease_owner == worker_id[:100],
                )
            )
        ).first()
        if job is None or job.lease_until is None or _aware(job.lease_until) <= _aware(utcnow()):
            raise MeetingExportError("EXPORT_LEASE_LOST", "export lease was lost")
        job.state = MeetingJobState.RUNNING.value
        job.updated_at = utcnow()
        payload = decode_json(job.input_json, {})
        try:
            request = MeetingExportCreateRequest.model_validate(
                {
                    "format": payload.get("format"),
                    "content": payload.get("content", "transcript"),
                    "transcript_source": payload.get("transcript_source", "display"),
                    "artifact_id": payload.get("source_artifact_id"),
                    "artifact_version": payload.get("source_artifact_version"),
                }
            )
        except Exception as exc:
            raise MeetingExportError("EXPORT_INPUT_INVALID", "export input is invalid") from exc
        meeting = await self.session.get(MeetingSession, job.meeting_id)
        if meeting is None or meeting.state == "deleted":
            raise MeetingExportError("EXPORT_MEETING_UNAVAILABLE", "meeting is unavailable")
        self.session.add(job)
        await self.session.commit()
        return meeting, request

    async def _render(
        self,
        meeting: MeetingSession,
        export_artifact: MeetingArtifact,
        job: MeetingJob,
        request: MeetingExportCreateRequest,
        owner_user_id: int,
    ) -> RenderedMeetingExport:
        export_format = str(request.format)
        if export_format == MeetingExportFormat.PDF.value:
            raise MeetingExportError(
                "EXPORT_FORMAT_NOT_AVAILABLE",
                f"{export_format.upper()} export is not available in this release",
            )
        segments = await self._transcript(
            meeting.id,
            owner_user_id,
            watermark=job.input_watermark,
        )
        revision_vector = {item.id: item.current_revision_no for item in segments}
        transcript_revision = sum(revision_vector.values())
        filename = _safe_filename(meeting.title, export_artifact.id, export_format)
        media_type = _MEDIA_TYPES[export_format]
        if request.content == MeetingExportContent.TRANSCRIPT:
            payload = _render_transcript(
                meeting,
                segments,
                export_format,
                str(request.transcript_source),
            )
            return RenderedMeetingExport(
                payload,
                filename,
                media_type,
                transcript_revision,
                revision_vector,
            )
        source_artifact = await self._minutes_artifact(
            meeting.id,
            owner_user_id,
            request.artifact_id,
            request.artifact_version,
        )
        payload = _render_minutes(
            meeting,
            source_artifact,
            segments,
            export_format,
            str(request.transcript_source),
        )
        return RenderedMeetingExport(
            payload,
            filename,
            media_type,
            transcript_revision,
            revision_vector,
            source_artifact.id,
            source_artifact.version,
        )

    async def _transcript(
        self,
        meeting_id: str,
        owner_user_id: int,
        *,
        watermark: int,
    ) -> list[TranscriptSegmentResponse]:
        values: list[TranscriptSegmentResponse] = []
        after = 0
        while len(values) < self.settings.max_segments:
            page, next_ordinal = await self.repository.transcript_page(
                meeting_id,
                owner_user_id,
                after_ordinal=after,
                limit=min(500, self.settings.max_segments - len(values)),
            )
            for segment in page:
                if segment.ordinal <= watermark:
                    values.append(segment)
            if next_ordinal is None or next_ordinal <= after or not page:
                break
            if page[-1].ordinal >= watermark:
                break
            after = next_ordinal
        if len(values) >= self.settings.max_segments and (values and values[-1].ordinal < watermark):
            raise MeetingExportError("EXPORT_TOO_LARGE", "transcript exceeds segment limit")
        return values

    async def _minutes_artifact(
        self,
        meeting_id: str,
        owner_user_id: int,
        artifact_id: str | None,
        artifact_version: int | None,
    ) -> MeetingArtifact:
        candidates: list[MeetingArtifact]
        if artifact_id:
            try:
                candidates = [await self.repository.get_artifact(meeting_id, artifact_id, owner_user_id)]
            except MeetingResourceNotFound as exc:
                raise MeetingExportError("EXPORT_ARTIFACT_NOT_FOUND", "minutes artifact was not found") from exc
        else:
            candidates = [
                value
                for value in await self.repository.list_artifacts(meeting_id, owner_user_id)
                if value.version == artifact_version
                and value.artifact_type
                in {
                    ArtifactType.ROLLING_MINUTES.value,
                    ArtifactType.FINAL_MINUTES.value,
                }
            ]
        candidates = [
            value
            for value in candidates
            if value.artifact_type
            in {
                ArtifactType.ROLLING_MINUTES.value,
                ArtifactType.FINAL_MINUTES.value,
            }
        ]
        if artifact_version is not None:
            candidates = [value for value in candidates if value.version == artifact_version]
        if not candidates:
            raise MeetingExportError("EXPORT_ARTIFACT_NOT_FOUND", "minutes artifact was not found")
        if len(candidates) > 1:
            raise MeetingExportError("EXPORT_ARTIFACT_AMBIGUOUS", "artifact version matches multiple minutes artifacts")
        artifact = candidates[0]
        if artifact.state not in {ArtifactState.READY.value, ArtifactState.STALE.value}:
            raise MeetingExportError("EXPORT_ARTIFACT_NOT_READY", "minutes artifact is not ready")
        return artifact

    async def _complete(
        self,
        job_id: str,
        worker_id: str,
        rendered: RenderedMeetingExport,
        persisted: PersistedMeetingExport,
    ) -> None:
        job = (
            await self.session.exec(
                select(MeetingJob).where(
                    MeetingJob.id == job_id,
                    MeetingJob.state == MeetingJobState.RUNNING.value,
                    MeetingJob.lease_owner == worker_id[:100],
                )
            )
        ).first()
        if job is None:
            raise MeetingExportError("EXPORT_LEASE_LOST", "export lease was lost")
        payload = decode_json(job.input_json, {})
        artifact = await self.session.get(MeetingArtifact, str(payload.get("export_id") or ""))
        if artifact is None or artifact.meeting_id != job.meeting_id:
            raise MeetingExportError("EXPORT_RECORD_UNAVAILABLE", "export record is unavailable")
        meeting = await self.session.get(MeetingSession, job.meeting_id)
        if meeting is None:
            raise MeetingExportError("EXPORT_MEETING_UNAVAILABLE", "meeting is unavailable")
        current = await self._transcript(
            job.meeting_id,
            meeting.owner_user_id,
            watermark=job.input_watermark,
        )
        current_vector = {item.id: item.current_revision_no for item in current}
        if current_vector != rendered.revision_vector:
            raise MeetingExportError(
                "EXPORT_TRANSCRIPT_CHANGED",
                "transcript changed while export was generated",
                retryable=True,
            )
        payload["source_artifact_id"] = rendered.source_artifact_id
        payload["source_artifact_version"] = rendered.source_artifact_version
        job.input_json = encode_json(payload)
        job.state = MeetingJobState.SUCCEEDED.value
        job.lease_owner = None
        job.lease_until = None
        job.public_error_code = None
        job.internal_diagnostic = None
        job.updated_at = utcnow()
        artifact.state = ArtifactState.READY.value
        artifact.transcript_revision = rendered.transcript_revision
        artifact.content_json = encode_json(
            {
                "schema_version": "siq.meeting.export_file.v1",
                "storage_key": persisted.storage_key,
                "filename": persisted.filename,
                "media_type": persisted.media_type,
                "byte_size": persisted.byte_size,
                "sha256": persisted.sha256,
                "source_artifact_id": rendered.source_artifact_id,
                "source_artifact_version": rendered.source_artifact_version,
            }
        )
        artifact.updated_at = utcnow()
        self.session.add(job)
        self.session.add(artifact)
        await self.events.append(
            job.meeting_id,
            "export.ready",
            {
                "export_id": artifact.id,
                "job_id": job.id,
                "format": payload.get("format"),
                "content": payload.get("content"),
                "byte_size": persisted.byte_size,
                "sha256": persisted.sha256,
                "source_artifact_id": rendered.source_artifact_id,
                "source_artifact_version": rendered.source_artifact_version,
            },
        )
        await self.session.commit()

    async def _fail(self, job_id: str, worker_id: str, exc: BaseException) -> None:
        job = (
            await self.session.exec(
                select(MeetingJob).where(
                    MeetingJob.id == job_id,
                    MeetingJob.lease_owner == worker_id[:100],
                    col(MeetingJob.state).in_([MeetingJobState.LEASED.value, MeetingJobState.RUNNING.value]),
                )
            )
        ).first()
        if job is None:
            return
        error = (
            exc
            if isinstance(exc, MeetingExportError)
            else MeetingExportError("EXPORT_FAILED", "export generation failed", retryable=True)
        )
        retryable = error.retryable and job.attempt < job.max_attempts
        job.state = MeetingJobState.RETRY_WAIT.value if retryable else MeetingJobState.FAILED.value
        job.lease_owner = None
        job.lease_until = None
        job.public_error_code = error.code
        job.internal_diagnostic = _diagnostic(exc)
        job.updated_at = utcnow()
        self.session.add(job)
        payload = decode_json(job.input_json, {})
        artifact = await self.session.get(MeetingArtifact, str(payload.get("export_id") or ""))
        if artifact is not None and not retryable:
            artifact.state = ArtifactState.FAILED.value
            artifact.updated_at = utcnow()
            self.session.add(artifact)
        await self.events.append(
            job.meeting_id,
            "export.retry_wait" if retryable else "export.failed",
            {
                "export_id": payload.get("export_id"),
                "job_id": job.id,
                "attempt": job.attempt,
                "max_attempts": job.max_attempts,
                "error_code": error.code,
            },
        )
        await self.session.commit()

    async def _owned_export_artifact(self, export_id: str, owner_user_id: int) -> MeetingArtifact:
        artifact = (
            await self.session.exec(
                select(MeetingArtifact)
                .join(MeetingSession, MeetingSession.id == MeetingArtifact.meeting_id)
                .where(
                    MeetingArtifact.id == export_id,
                    MeetingArtifact.artifact_type == ArtifactType.EXPORT.value,
                    MeetingSession.owner_user_id == owner_user_id,
                    MeetingSession.state != "deleted",
                )
            )
        ).first()
        if artifact is None:
            raise MeetingResourceNotFound("meeting resource not found")
        return artifact

    async def list_exports(self, meeting_id: str, owner_user_id: int) -> list[MeetingExportResponse]:
        await self.repository.get_session(meeting_id, owner_user_id)
        artifacts = list(
            (
                await self.session.exec(
                    select(MeetingArtifact)
                    .where(
                        MeetingArtifact.meeting_id == meeting_id,
                        MeetingArtifact.artifact_type == ArtifactType.EXPORT.value,
                    )
                    .order_by(col(MeetingArtifact.created_at).desc())
                )
            ).all()
        )
        jobs = list(
            (
                await self.session.exec(
                    select(MeetingJob).where(
                        MeetingJob.meeting_id == meeting_id,
                        MeetingJob.job_kind == MeetingJobKind.EXPORT.value,
                    )
                )
            ).all()
        )
        by_export = {str(decode_json(job.input_json, {}).get("export_id") or ""): job for job in jobs}
        return [export_response(artifact, by_export[artifact.id]) for artifact in artifacts if artifact.id in by_export]

    async def issue_ticket(
        self,
        meeting_id: str,
        export_id: str,
        owner_user_id: int,
        *,
        origin: str,
    ) -> tuple[str, datetime, DownloadableMeetingExport]:
        artifact, job = await self.repository.get_export(meeting_id, export_id, owner_user_id)
        if artifact.state != ArtifactState.READY.value or job.state != MeetingJobState.SUCCEEDED.value:
            raise MeetingInvalidOperation("meeting export is not ready")
        metadata = decode_json(artifact.content_json, {})
        downloadable = self.storage.resolve(metadata)
        raw = secrets.token_urlsafe(32)
        expires_at = utcnow() + timedelta(seconds=self.settings.ticket_ttl_seconds)
        ticket = MeetingStreamTicket(
            token_hash=hashlib.sha256(raw.encode("ascii")).hexdigest(),
            meeting_id=meeting_id,
            owner_user_id=owner_user_id,
            stream_epoch=1,
            purpose="meeting_export_download",
            origin=origin[:500],
            expires_at=expires_at,
            connection_id=export_id,
        )
        self.session.add(ticket)
        await self.events.append(
            meeting_id,
            "export.download_ticket.issued",
            {"export_id": export_id, "expires_at": expires_at.isoformat()},
        )
        await self.session.commit()
        return raw, expires_at, downloadable

    async def consume_ticket(
        self,
        meeting_id: str,
        export_id: str,
        owner_user_id: int,
        raw_ticket: str,
    ) -> DownloadableMeetingExport:
        try:
            encoded_ticket = raw_ticket.encode("ascii")
        except UnicodeEncodeError as exc:
            raise MeetingResourceNotFound("meeting resource not found") from exc
        token_hash = hashlib.sha256(encoded_ticket).hexdigest()
        ticket = (
            await self.session.exec(
                select(MeetingStreamTicket).where(
                    MeetingStreamTicket.meeting_id == meeting_id,
                    MeetingStreamTicket.owner_user_id == owner_user_id,
                    MeetingStreamTicket.connection_id == export_id,
                    MeetingStreamTicket.purpose == "meeting_export_download",
                    MeetingStreamTicket.token_hash == token_hash,
                    MeetingStreamTicket.consumed_at.is_(None),
                )
            )
        ).first()
        if ticket is None or _aware(ticket.expires_at) <= _aware(utcnow()):
            raise MeetingResourceNotFound("meeting resource not found")
        artifact, job = await self.repository.get_export(meeting_id, export_id, owner_user_id)
        if artifact.state != ArtifactState.READY.value or job.state != MeetingJobState.SUCCEEDED.value:
            raise MeetingResourceNotFound("meeting resource not found")
        downloadable = self.storage.resolve(decode_json(artifact.content_json, {}))
        now = utcnow()
        result = await self.session.exec(
            update(MeetingStreamTicket)
            .where(
                MeetingStreamTicket.id == ticket.id,
                MeetingStreamTicket.consumed_at.is_(None),
            )
            .values(consumed_at=now)
        )
        if not result.rowcount:
            await self.session.rollback()
            raise MeetingResourceNotFound("meeting resource not found")
        await self.events.append(
            meeting_id,
            "export.downloaded",
            {
                "export_id": export_id,
                "byte_size": downloadable.byte_size,
                "sha256": downloadable.sha256,
            },
        )
        await self.session.commit()
        return downloadable


__all__ = [
    "DownloadableMeetingExport",
    "MeetingExportError",
    "MeetingExportService",
    "MeetingExportSettings",
    "MeetingExportStorage",
    "PersistedMeetingExport",
    "RenderedMeetingExport",
]
