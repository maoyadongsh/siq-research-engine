"""Local parsing providers that do not require an external service."""

from __future__ import annotations

import html
import json
import re
import zipfile
from pathlib import Path
from typing import Any

from contracts import ParseConfig, ParseOutput, SourceFile
from page_ranges import selected_page_indexes


HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")
HTML_TAG_RE = re.compile(r"<[^>]+>")
HTML_BLOCK_BREAK_RE = re.compile(r"</(?:p|div|section|article|h[1-6]|li|tr|table)>", re.I)


def _decode_bytes(path: Path) -> str:
    data = path.read_bytes()
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    return "\n".join(line.rstrip() for line in text.split("\n")).strip()


def _text_to_blocks(task_id: str, text: str, kind: str = "text") -> list[dict[str, Any]]:
    blocks: list[dict[str, Any]] = []
    paragraphs: list[str] = []
    order = 1

    def flush_paragraph() -> None:
        nonlocal order
        if not paragraphs:
            return
        paragraph = "\n".join(paragraphs).strip()
        paragraphs.clear()
        if not paragraph:
            return
        block_id = f"b{order:06d}"
        blocks.append(
            {
                "block_id": block_id,
                "type": "paragraph",
                "sub_type": kind,
                "text": paragraph,
                "markdown": paragraph,
                "html": "",
                "page_number": 1,
                "page_index": 0,
                "sheet_name": "",
                "slide_number": None,
                "bbox": [],
                "bbox_unit": "none",
                "reading_order": order,
                "parent_block_id": "",
                "source_ref": {
                    "evidence_id": f"doc:{task_id}:p1:{block_id}",
                    "source_type": f"{kind}_block",
                    "path": "",
                },
                "confidence": 1.0,
                "warnings": [],
            }
        )
        order += 1

    for raw_line in text.split("\n"):
        line = raw_line.strip()
        if not line:
            flush_paragraph()
            continue
        heading = HEADING_RE.match(line)
        if heading:
            flush_paragraph()
            block_id = f"b{order:06d}"
            level = len(heading.group(1))
            title = heading.group(2).strip()
            blocks.append(
                {
                    "block_id": block_id,
                    "type": "title" if level == 1 else "heading",
                    "sub_type": f"h{level}",
                    "text": title,
                    "markdown": line,
                    "html": "",
                    "page_number": 1,
                    "page_index": 0,
                    "sheet_name": "",
                    "slide_number": None,
                    "bbox": [],
                    "bbox_unit": "none",
                    "reading_order": order,
                    "parent_block_id": "",
                    "source_ref": {
                        "evidence_id": f"doc:{task_id}:p1:{block_id}",
                        "source_type": f"{kind}_heading",
                        "path": "",
                    },
                    "confidence": 1.0,
                    "warnings": [],
                }
            )
            order += 1
        else:
            paragraphs.append(line)
    flush_paragraph()
    return blocks


def _blocks_to_markdown(blocks: list[dict[str, Any]]) -> str:
    parts = []
    for block in blocks:
        block_id = block.get("block_id", "")
        evidence_id = (block.get("source_ref") or {}).get("evidence_id", "")
        page = block.get("page_number") or 1
        marker = f"<!-- DOC_BLOCK: {block_id} page={page} evidence={evidence_id} -->"
        markdown = str(block.get("markdown") or block.get("text") or "").strip()
        if markdown:
            parts.append(f"{marker}\n{markdown}")
    return "\n\n".join(parts).strip() + ("\n" if parts else "")


def parse_text_document(task_id: str, source: SourceFile, config: ParseConfig) -> ParseOutput:
    text = _normalize_text(_decode_bytes(source.path))
    blocks = _text_to_blocks(task_id, text, "text")
    warnings = []
    if not blocks:
        warnings.append(
            {
                "code": "empty_text",
                "severity": "warning",
                "message": "文档没有可解析文本内容。",
            }
        )
    return ParseOutput(
        markdown=_blocks_to_markdown(blocks),
        blocks=blocks,
        warnings=warnings,
        page_count=1,
        provider_name="simple_text_parser",
        document_kind="text",
    )


def parse_html_document(task_id: str, source: SourceFile, config: ParseConfig) -> ParseOutput:
    raw = _decode_bytes(source.path)
    raw = re.sub(r"(?is)<(script|style).*?</\1>", " ", raw)
    raw = HTML_BLOCK_BREAK_RE.sub("\n", raw)
    text = HTML_TAG_RE.sub(" ", raw)
    text = html.unescape(text)
    text = re.sub(r"[ \t\f\v]+", " ", text)
    text = _normalize_text(text)
    blocks = _text_to_blocks(task_id, text, "html")
    return ParseOutput(
        markdown=_blocks_to_markdown(blocks),
        blocks=blocks,
        warnings=[] if blocks else [{"code": "empty_html", "severity": "warning", "message": "HTML 正文提取为空。"}],
        page_count=1,
        provider_name="html_reader",
        document_kind="html",
    )


def parse_pdf_document(task_id: str, source: SourceFile, config: ParseConfig) -> ParseOutput:
    warnings: list[dict[str, Any]] = []
    pages: list[tuple[int, str]] = []
    provider = "pypdf_text_parser"
    try:
        from pypdf import PdfReader  # type: ignore

        reader = PdfReader(str(source.path))
        for page_index in selected_page_indexes(config.page_ranges, len(reader.pages)):
            pages.append((page_index + 1, _normalize_text(reader.pages[page_index].extract_text() or "")))
    except ValueError:
        raise
    except Exception as exc:  # pragma: no cover - depends on optional parser behavior
        provider = "pdf_metadata_parser"
        warnings.append(
            {
                "code": "pdf_text_extract_unavailable",
                "severity": "warning",
                "message": f"PDF 文本提取不可用，已生成占位产物: {exc}",
            }
        )

    if not pages:
        pages = [(1, "")]

    blocks: list[dict[str, Any]] = []
    order = 1
    for fallback_index, (page_number, page_text) in enumerate(pages):
        page_index = page_number - 1 if page_number > 0 else fallback_index
        page_blocks = _text_to_blocks(task_id, page_text, "pdf")
        if not page_blocks:
            block_id = f"b{order:06d}"
            page_blocks = [
                {
                    "block_id": block_id,
                    "type": "unknown",
                    "sub_type": "pdf_page_placeholder",
                    "text": "",
                    "markdown": f"<!-- 第 {page_number} 页暂无可提取文本 -->",
                    "html": "",
                    "page_number": page_number,
                    "page_index": page_index,
                    "sheet_name": "",
                    "slide_number": None,
                    "bbox": [],
                    "bbox_unit": "none",
                    "reading_order": order,
                    "parent_block_id": "",
                    "source_ref": {
                        "evidence_id": f"doc:{task_id}:p{page_number}:{block_id}",
                        "source_type": "pdf_page",
                        "path": f"raw/original/{source.filename}",
                    },
                    "confidence": 0.0,
                    "warnings": ["no_text"],
                }
            ]
        for block in page_blocks:
            block["block_id"] = f"b{order:06d}"
            block["page_number"] = page_number
            block["page_index"] = page_index
            block["reading_order"] = order
            block["source_ref"] = {
                "evidence_id": f"doc:{task_id}:p{page_number}:b{order:06d}",
                "source_type": "pdf_block",
                "path": f"raw/original/{source.filename}",
            }
            order += 1
            blocks.append(block)

    empty_pages = sum(1 for _, page in pages if not page.strip())
    if empty_pages:
        warnings.append(
            {
                "code": "pdf_pages_without_text",
                "severity": "warning",
                "message": f"{empty_pages} 页未提取到文本，可能需要 OCR 或 MinerU provider。",
            }
        )
    return ParseOutput(
        markdown=_blocks_to_markdown(blocks),
        blocks=blocks,
        warnings=warnings,
        page_count=len(pages),
        provider_name=provider,
        document_kind="pdf",
    )


def parse_docx_document(task_id: str, source: SourceFile, config: ParseConfig) -> ParseOutput:
    warnings: list[dict[str, Any]] = []
    text = ""
    try:
        import docx  # type: ignore

        doc = docx.Document(str(source.path))
        text = "\n\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            if rows:
                tables.append(rows)
        if tables:
            text += "\n\n" + "\n\n".join(_markdown_table(rows) for rows in tables)
    except Exception as exc:
        warnings.append(
            {
                "code": "docx_parser_unavailable",
                "severity": "warning",
                "message": f"DOCX 解析依赖不可用或文件无法读取: {exc}",
            }
        )
        text = f"# {source.filename}\n\n此 Office 文档已入队归档，但当前运行环境未启用 DOCX 深度解析。"
    blocks = _text_to_blocks(task_id, _normalize_text(text), "word")
    return ParseOutput(
        markdown=_blocks_to_markdown(blocks),
        blocks=blocks,
        warnings=warnings,
        page_count=1,
        provider_name="office_local",
        document_kind="word",
    )


def parse_spreadsheet_document(task_id: str, source: SourceFile, config: ParseConfig) -> ParseOutput:
    warnings: list[dict[str, Any]] = []
    markdown_parts: list[str] = []
    physical_tables: list[dict[str, Any]] = []
    try:
        import openpyxl  # type: ignore

        workbook = openpyxl.load_workbook(str(source.path), read_only=True, data_only=True)
        table_order = 1
        for sheet in workbook.worksheets:
            rows = []
            for row in sheet.iter_rows(values_only=True):
                values = ["" if value is None else str(value) for value in row]
                if any(value.strip() for value in values):
                    rows.append(values)
                if len(rows) >= 80:
                    break
            if not rows:
                continue
            markdown = _markdown_table(rows[:40])
            markdown_parts.append(f"## {sheet.title}\n\n{markdown}")
            table_id = f"pt-{table_order:06d}"
            physical_tables.append(
                {
                    "table_id": table_id,
                    "block_id": "",
                    "title": sheet.title,
                    "caption": sheet.title,
                    "page_number": 1,
                    "sheet_name": sheet.title,
                    "html": "",
                    "markdown": markdown,
                    "cells": _table_cells(task_id, table_id, rows, sheet.title),
                    "quality": {
                        "has_header": bool(rows),
                        "row_count": len(rows),
                        "column_count": max((len(row) for row in rows), default=0),
                        "empty_cell_ratio": _empty_cell_ratio(rows),
                    },
                }
            )
            table_order += 1
    except Exception as exc:
        warnings.append(
            {
                "code": "spreadsheet_parser_unavailable",
                "severity": "warning",
                "message": f"Excel 解析依赖不可用或文件无法读取: {exc}",
            }
        )
        markdown_parts.append(f"# {source.filename}\n\n此电子表格已归档，但当前运行环境未启用表格深度解析。")

    text = "\n\n".join(markdown_parts)
    blocks = _text_to_blocks(task_id, text, "excel")
    for index, table in enumerate(physical_tables):
        if index < len(blocks):
            table["block_id"] = blocks[index]["block_id"]
    return ParseOutput(
        markdown=_blocks_to_markdown(blocks),
        blocks=blocks,
        tables=physical_tables,
        warnings=warnings,
        page_count=1,
        provider_name="spreadsheet_parser",
        document_kind="excel",
    )


def parse_image_document(task_id: str, source: SourceFile, config: ParseConfig) -> ParseOutput:
    block_id = "b000001"
    image_id = "img-000001"
    evidence_id = f"doc:{task_id}:p1:{image_id}"
    markdown = f'<a id="md-{image_id}"></a>\n\n![{source.filename}](images/original/{source.filename})\n'
    block = {
        "block_id": block_id,
        "type": "image",
        "sub_type": "original_image",
        "text": source.filename,
        "markdown": markdown,
        "html": "",
        "page_number": 1,
        "page_index": 0,
        "sheet_name": "",
        "slide_number": None,
        "bbox": [],
        "bbox_unit": "none",
        "reading_order": 1,
        "parent_block_id": "",
        "source_ref": {
            "evidence_id": evidence_id,
            "source_type": "image_file",
            "path": f"images/original/{source.filename}",
        },
        "confidence": 1.0,
        "warnings": ["ocr_not_enabled"],
    }
    figure = {
        "image_id": image_id,
        "block_id": block_id,
        "type": "image",
        "page_number": 1,
        "page_index": 0,
        "bbox": [],
        "bbox_unit": "none",
        "image_path": f"images/original/{source.filename}",
        "thumbnail_path": "",
        "source_page_image_path": f"images/original/{source.filename}",
        "caption": source.filename,
        "footnote": "",
        "nearby_heading": "",
        "ocr_text": "",
        "alt_text": source.filename,
        "markdown": markdown,
        "markdown_anchor": f"md-{image_id}",
        "evidence_id": evidence_id,
        "quality": {
            "crop_available": True,
            "caption_detected": False,
            "ocr_available": False,
            "is_low_resolution": False,
        },
    }
    return ParseOutput(
        markdown=f"<!-- DOC_BLOCK: {block_id} page=1 evidence={evidence_id} -->\n{markdown}",
        blocks=[block],
        figures=[figure],
        warnings=[{"code": "image_ocr_not_enabled", "severity": "warning", "message": "图片已归档，当前 P0 provider 未执行 OCR。"}],
        page_count=1,
        provider_name="image_local",
        document_kind="image",
    )


def parse_office_placeholder(task_id: str, source: SourceFile, config: ParseConfig, kind: str) -> ParseOutput:
    text = f"# {source.filename}\n\n此 {kind.upper()} 文档已归档。当前本地 provider 生成占位产物，后续可切换 MinerU/LibreOffice provider 做深度解析。"
    blocks = _text_to_blocks(task_id, text, kind)
    return ParseOutput(
        markdown=_blocks_to_markdown(blocks),
        blocks=blocks,
        warnings=[
            {
                "code": f"{kind}_deep_parse_not_enabled",
                "severity": "warning",
                "message": f"{kind.upper()} 深度解析 provider 未启用。",
            }
        ],
        page_count=1,
        provider_name="office_placeholder",
        document_kind=kind,
    )


def parse_json_schema_excerpt(schema: dict[str, Any], markdown: str) -> dict[str, Any]:
    properties = schema.get("properties") if isinstance(schema, dict) else {}
    if not isinstance(properties, dict):
        properties = {}
    result: dict[str, Any] = {}
    text = markdown or ""
    for key in properties:
        pattern = re.compile(rf"(?im)^\s*{re.escape(str(key))}\s*[:：]\s*(.+?)\s*$")
        match = pattern.search(text)
        result[key] = match.group(1).strip() if match else None
    return result


def _markdown_table(rows: list[list[Any]]) -> str:
    if not rows:
        return ""
    width = max((len(row) for row in rows), default=0)
    normalized = [[str(cell or "").replace("\n", " ").strip() for cell in row] + [""] * (width - len(row)) for row in rows]
    header = normalized[0]
    body = normalized[1:] or [[""] * width]
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join("---" for _ in range(width)) + " |",
    ]
    lines.extend("| " + " | ".join(row) + " |" for row in body)
    return "\n".join(lines)


def _table_cells(task_id: str, table_id: str, rows: list[list[Any]], sheet_name: str = "") -> list[dict[str, Any]]:
    cells = []
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            text = str(value or "")
            cells.append(
                {
                    "row_index": row_index,
                    "column_index": column_index,
                    "text": text,
                    "bbox": [],
                    "evidence_id": f"doc:{task_id}:sheet:{sheet_name}:{table_id}:r{row_index}:c{column_index}",
                }
            )
    return cells


def _empty_cell_ratio(rows: list[list[Any]]) -> float:
    total = 0
    empty = 0
    for row in rows:
        for value in row:
            total += 1
            if not str(value or "").strip():
                empty += 1
    return round(empty / total, 4) if total else 0.0


def docx_xml_text(path: Path) -> str:
    with zipfile.ZipFile(path) as archive:
        xml = archive.read("word/document.xml").decode("utf-8", errors="replace")
    return _normalize_text(HTML_TAG_RE.sub(" ", xml))
